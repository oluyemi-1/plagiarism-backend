# services/text_extractor.py - Phase 2: Text Extraction Service
import os
import logging
from typing import Dict, Any, Optional
import chardet
from PyPDF2 import PdfReader
from docx import Document
import io

logger = logging.getLogger(__name__)

class TextExtractor:
    """Service for extracting text from various file formats"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.txt'}
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return list(self.supported_formats)
    
    async def extract_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from uploaded file
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename with extension
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract text based on file type
            if file_extension == '.pdf':
                extracted_text = await self._extract_from_pdf(file_content)
            elif file_extension == '.docx':
                extracted_text = await self._extract_from_docx(file_content)
            elif file_extension == '.txt':
                extracted_text = await self._extract_from_txt(file_content)
            else:
                raise ValueError(f"Handler not implemented for: {file_extension}")
            
            # Calculate basic statistics
            word_count = len(extracted_text.split()) if extracted_text else 0
            char_count = len(extracted_text) if extracted_text else 0
            
            return {
                "text": extracted_text,
                "word_count": word_count,
                "character_count": char_count,
                "file_type": file_extension,
                "filename": filename,
                "success": True,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}")
            return {
                "text": "",
                "word_count": 0,
                "character_count": 0,
                "file_type": file_extension if 'file_extension' in locals() else "unknown",
                "filename": filename,
                "success": False,
                "error": str(e)
            }
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            # Create a file-like object from bytes
            pdf_file = io.BytesIO(file_content)
            
            # Read PDF
            reader = PdfReader(pdf_file)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            extracted_text = "\n\n".join(text_parts)
            
            if not extracted_text.strip():
                raise ValueError("No readable text found in PDF")
            
            return extracted_text.strip()
            
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            # Create a file-like object from bytes
            docx_file = io.BytesIO(file_content)
            
            # Read DOCX
            document = Document(docx_file)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract text from tables if present
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            extracted_text = "\n\n".join(text_parts)
            
            if not extracted_text.strip():
                raise ValueError("No readable text found in DOCX")
            
            return extracted_text.strip()
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    async def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Detect encoding
            detected = chardet.detect(file_content)
            encoding = detected.get('encoding', 'utf-8')
            
            # If confidence is too low, try common encodings
            if detected.get('confidence', 0) < 0.7:
                encodings_to_try = ['utf-8', 'ascii', 'latin1', 'cp1252']
                
                for enc in encodings_to_try:
                    try:
                        text = file_content.decode(enc)
                        return text.strip()
                    except UnicodeDecodeError:
                        continue
                
                # If all fail, use utf-8 with error handling
                text = file_content.decode('utf-8', errors='ignore')
            else:
                text = file_content.decode(encoding)
            
            if not text.strip():
                raise ValueError("No readable text found in file")
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Text file extraction failed: {str(e)}")